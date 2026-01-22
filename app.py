import streamlit as st
import google.generativeai as genai
import sqlite3
import uuid
import os
import tempfile
import re
from datetime import datetime
import time
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- CONFIGURARE PAGINĂ ---
st.set_page_config(page_title="Consultant Vânzări IT AI", layout="wide")

# --- 1. GESTIONARE BAZĂ DE DATE (SQLite) ---
DB_FILE = "chat_history.db"

def init_db():
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS messages
                 (session_id TEXT, role TEXT, content TEXT, timestamp DATETIME)''')
    conn.commit()
    conn.close()

def save_message(session_id, role, content):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("INSERT INTO messages VALUES (?, ?, ?, ?)", 
              (session_id, role, content, datetime.now()))
    conn.commit()
    conn.close()

def load_history(session_id):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("SELECT role, content FROM messages WHERE session_id = ? ORDER BY timestamp", (session_id,))
    rows = c.fetchall()
    conn.close()
    history = []
    for role, content in rows:
        history.append({"role": role, "content": content})
    return history

def clear_session_history(session_id):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("DELETE FROM messages WHERE session_id = ?", (session_id,))
    conn.commit()
    conn.close()

init_db()

# --- 2. GESTIONARE ID SESIUNE ---
if "session_id" not in st.query_params:
    new_id = str(uuid.uuid4())
    st.query_params["session_id"] = new_id
    session_id = new_id
else:
    session_id = st.query_params["session_id"]

# --- 3. GESTIONARE CHEI API ---
def configure_gemini():
    # Căutăm cheile în secrets sub numele 'GOOGLE_API_KEYS' sau 'api_keys'
    api_keys = []
    
    # Verificăm ambele denumiri posibile în secrets
    if "GOOGLE_API_KEYS" in st.secrets:
        if isinstance(st.secrets["GOOGLE_API_KEYS"], list):
            api_keys = st.secrets["GOOGLE_API_KEYS"]
        else:
            api_keys = st.secrets["GOOGLE_API_KEYS"].split(",")
    elif "api_keys" in st.secrets:
        if isinstance(st.secrets["api_keys"], list):
            api_keys = st.secrets["api_keys"]
        else:
            api_keys = st.secrets["api_keys"].split(",")
    
    valid_model = None
    
    # Rotație chei
    for key in api_keys:
        key = key.strip()
        try:
            genai.configure(api_key=key)
            # CORECTAT: Folosim gemini-2.5-flash
            model = genai.GenerativeModel('gemini-2.5-flash')
            # Testăm conexiunea
            model.generate_content("test", request_options={"timeout": 5})
            valid_model = model
            break 
        except Exception:
            continue

    # Fallback: Cheie manuală în browser
    if not valid_model:
        st.sidebar.warning("Nicio cheie din sistem nu merge. Introdu una manual.")
        user_key = st.sidebar.text_input("Cheie API Gemini:", type="password")
        if user_key:
            try:
                genai.configure(api_key=user_key)
                model = genai.GenerativeModel('gemini-2.5-flash')
                valid_model = model
            except:
                pass
    
    return valid_model

# --- 4. FUNCȚII UPLOAD ---
def upload_to_gemini(uploaded_file):
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_path = tmp_file.name
        
        # Upload
        gemini_file = genai.upload_file(path=tmp_path, display_name=uploaded_file.name)
        
        # Așteptare procesare
        while gemini_file.state.name == "PROCESSING":
            time.sleep(1)
            gemini_file = genai.get_file(gemini_file.name)
            
        os.remove(tmp_path)
        return gemini_file
    except Exception as e:
        st.error(f"Eroare upload: {e}")
        return None

# --- 5. GENERATOR DOCUMENT WORD ---
def add_markdown_paragraph(doc, text):
    p = doc.add_paragraph()
    # Gestionare Bold (**text**)
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)

def create_docx(markdown_text):
    doc = Document()
    doc.add_heading('Ofertă / Raport AI', 0)

    lines = markdown_text.split('\n')
    table_buffer = [] 
    
    for line in lines:
        line = line.strip()
        
        # Detectare tabel Markdown
        if line.startswith('|') and line.endswith('|'):
            if '---' in line: 
                continue 
            cells = [c.strip() for c in line.split('|')[1:-1]]
            table_buffer.append(cells)
        else:
            # Scriem tabelul acumulat anterior
            if table_buffer:
                if len(table_buffer) > 0:
                    rows = len(table_buffer)
                    cols = len(table_buffer[0])
                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = 'Table Grid'
                    for i, row_data in enumerate(table_buffer):
                        row_cells = table.rows[i].cells
                        for j, cell_text in enumerate(row_data):
                            if j < len(row_cells):
                                row_cells[j].text = cell_text
                                # Bold pe header
                                if i == 0:
                                    for paragraph in row_cells[j].paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True
                table_buffer = [] 
                doc.add_paragraph() 

            # Procesare text normal
            if line:
                if line.startswith('###'):
                    doc.add_heading(line.replace('###', '').strip(), level=3)
                elif line.startswith('##'):
                    doc.add_heading(line.replace('##', '').strip(), level=2)
                elif line.startswith('#'):
                    doc.add_heading(line.replace('#', '').strip(), level=1)
                elif line.startswith('- '):
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(line[2:])
                else:
                    add_markdown_paragraph(doc, line)

    # Dacă a rămas un tabel la final
    if table_buffer:
        rows = len(table_buffer)
        cols = len(table_buffer[0])
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        for i, row_data in enumerate(table_buffer):
            row_cells = table.rows[i].cells
            for j, cell_text in enumerate(row_data):
                if j < len(row_cells):
                    row_cells[j].text = cell_text
                    if i == 0:
                        for paragraph in row_cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- INTERFAȚA GRAFICĂ (UI) ---

st.title("🤖 Consultant Vânzări IT - AI")
st.markdown(f"**ID Sesiune:** `{session_id}` (Salvează acest link pentru a reveni)")

model = configure_gemini()

with st.sidebar:
    st.header("📂 Documente Companie")
    portfolio_file = st.file_uploader("Portofoliu Companie (PDF)", type=['pdf'], key="port")
    catalog_file = st.file_uploader("Catalog Produse & Prețuri (PDF/CSV)", type=['pdf', 'txt', 'csv'], key="cat")
    
    st.divider()
    st.header("📋 Documente Client")
    st.info("Încarcă lista de necesar primită de la client.")
    client_req_file = st.file_uploader("Cerințe Client (PDF/CSV/TXT)", type=['pdf', 'txt', 'csv'], key="req")
    
    if st.button("Procesează Toate Documentele"):
        if model:
            with st.spinner("Se analizează fișierele pe serverele Google..."):
                # 1. Portofoliu
                if portfolio_file:
                    f1 = upload_to_gemini(portfolio_file)
                    if f1: 
                        st.session_state['portfolio_ref'] = f1
                        st.success("✅ Portofoliu Procesat")
                
                # 2. Catalog
                if catalog_file:
                    f2 = upload_to_gemini(catalog_file)
                    if f2: 
                        st.session_state['catalog_ref'] = f2
                        st.success("✅ Catalog Procesat")

                # 3. Cerințe Client
                if client_req_file:
                    f3 = upload_to_gemini(client_req_file)
                    if f3:
                        st.session_state['client_req_ref'] = f3
                        st.success("✅ Cerințe Client Procesate")
        else:
            st.error("Modelul AI nu este configurat. Verifică cheile API.")

    st.divider()
    if st.button("RESET CONVERSAȚIE & FIȘIERE", type="primary"):
        clear_session_history(session_id)
        # Resetăm și fișierele din sesiune
        keys_to_remove = ['portfolio_ref', 'catalog_ref', 'client_req_ref']
        for key in keys_to_remove:
            if key in st.session_state:
                del st.session_state[key]
        st.rerun()

# Recuperare istoric din SQLite
if "messages" not in st.session_state:
    st.session_state.messages = load_history(session_id)

for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# Input utilizator
if prompt := st.chat_input("Ex: Analizează fișierul clientului și fă o ofertă din catalog..."):
    if not model:
        st.error("Te rog configurează o cheie API validă.")
    else:
        # 1. Salvare user input
        st.session_state.messages.append({"role": "user", "content": prompt})
        save_message(session_id, "user", prompt)
        with st.chat_message("user"):
            st.markdown(prompt)

        # 2. Construire Context
        system_instruction = """
        Ești un consultant expert în vânzări de echipamente IT.
        Sarcina ta:
        1. Analizează documentele încărcate.
        2. Prioritate MAXIMĂ: Dacă există 'Cerințe Client', identifică produsele cerute.
        3. Caută produsele identificate în 'Catalog'. 
           - Dacă există exact: folosește prețul din catalog.
           - Dacă nu există exact: propune cea mai bună alternativă din catalog și explică de ce.
        4. Generează oferta finală sub formă de tabel Markdown detaliat.
        5. Fii politicos și profesionist.
        """
        
        current_request = [system_instruction]
        
        if 'portfolio_ref' in st.session_state:
            current_request.append("DOCUMENT CONTEXT: Portofoliu Companie")
            current_request.append(st.session_state['portfolio_ref'])
            
        if 'catalog_ref' in st.session_state:
            current_request.append("DOCUMENT CONTEXT: Catalog Produse și Prețuri")
            current_request.append(st.session_state['catalog_ref'])

        if 'client_req_ref' in st.session_state:
            current_request.append("DOCUMENT CRITIC: Lista de Cerințe a Clientului")
            current_request.append(st.session_state['client_req_ref'])
            
        # Adăugăm ultimele 5 mesaje pentru context
        history_text = "\n".join([f"{m['role'].upper()}: {m['content']}" for m in st.session_state.messages[-5:]])
        current_request.append(f"Istoric Discuție:\n{history_text}")
        current_request.append(f"SOLICITARE CURENTĂ: {prompt}")

        # 3. Generare Răspuns
        with st.chat_message("assistant"):
            with st.spinner("Generez oferta comparativă..."):
                try:
                    response = model.generate_content(current_request)
                    response_text = response.text
                    
                    st.markdown(response_text)
                    
                    # Salvare
                    st.session_state.messages.append({"role": "assistant", "content": response_text})
                    save_message(session_id, "assistant", response_text)

                    # Generare Word
                    docx_file = create_docx(response_text)
                    
                    st.download_button(
                        label="📄 Descarcă Oferta (Word .docx)",
                        data=docx_file,
                        file_name=f"Oferta_Generata_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                except Exception as e:
                    st.error(f"Eroare la generare: {e}")
